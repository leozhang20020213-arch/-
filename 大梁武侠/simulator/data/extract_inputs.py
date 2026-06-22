from __future__ import annotations

import json
import re
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Any

from docx import Document
from openpyxl import load_workbook


ROOT = Path(__file__).resolve().parents[2]
REPORTS = ROOT / "reports"
RAW_LOGS = REPORTS / "raw_logs"


RULE_KEYWORDS = [
    "检定",
    "掷骰",
    "得利",
    "破绽",
    "势态",
    "Pos",
    "危机",
    "解密",
    "名望",
    "人情",
    "情报",
    "把柄",
    "休整",
    "岁月",
    "境界",
    "感悟",
    "武学",
    "内功",
    "轻功",
    "护甲",
    "奇物",
    "敌人",
]

TERMS = [
    "得利",
    "破绽",
    "势态",
    "Pos",
    "名望",
    "人情",
    "情报",
    "把柄",
    "场景进度",
    "危机",
    "主修",
    "兼修",
    "旁学",
    "境界",
    "感悟",
    "休整",
    "岁月",
    "江湖事件",
]


@dataclass
class SourceFile:
    expected_name: str
    found_path: str | None
    status: str
    note: str = ""


@dataclass
class RuleRecord:
    rule_id: str
    rule_name: str
    source: str
    chapter: str
    category: str
    applicable_scene: str
    input_conditions: str
    resolution: str
    affected_resources: str
    linked_database: str
    needs_dm_judgment: str
    ambiguity: str
    excerpt: str


@dataclass
class DatabaseRecord:
    workbook: str
    sheet_name: str
    primary_key_guess: str
    row_count: int
    column_count: int
    core_fields: str
    called_by_rules: str
    fields_without_explanation: str
    possibly_unused_fields: str


@dataclass
class TermRecord:
    term: str
    source: str
    occurrences: int
    first_chapter: str
    first_context: str
    needs_definition_check: str


def find_inputs() -> list[SourceFile]:
    expected = [
        "大梁武侠TRPG_玩家规则书_完整整合版_v0.41.docx",
        "大梁武侠TRPG_DM规则书_完整整合版_v0.41.docx",
        "大梁武侠TRPG_数据库填充_V1.xlsx",
        "大梁武侠TRPG_岁月与生成数据库_V1.xlsx",
        "大梁武侠TRPG_世界投放与剧本生成数据库_V1.xlsx",
        "大梁武侠TRPG_总数据库索引与规则书成稿计划_V1.xlsx",
    ]
    all_files = [p for p in ROOT.rglob("*") if p.is_file() and not p.name.startswith("~$")]
    found: list[SourceFile] = []
    for name in expected:
        exact = [p for p in all_files if p.name == name]
        if exact:
            found.append(SourceFile(name, str(exact[0].relative_to(ROOT)), "found"))
            continue
        distinctive_tokens = [t for t in ["玩家规则书", "DM规则书", "数据库填充", "岁月", "世界投放", "总数据库索引"] if t in name]
        tokens = [t for t in re.split(r"[_\s]+", name.replace(".docx", "").replace(".xlsx", "")) if t]
        candidates = []
        for p in all_files:
            if distinctive_tokens and not any(t in p.name for t in distinctive_tokens):
                continue
            score = sum(1 for t in tokens if t and t in p.name)
            if score >= max(2, len(tokens) // 2):
                candidates.append((score, p))
        candidates.sort(reverse=True, key=lambda x: x[0])
        if candidates:
            found.append(SourceFile(name, str(candidates[0][1].relative_to(ROOT)), "near_match", f"score={candidates[0][0]}"))
        else:
            found.append(SourceFile(name, None, "missing", "项目目录内未找到同名或近似文件"))
    return found


def docx_blocks(path: Path) -> list[dict[str, str]]:
    doc = Document(str(path))
    blocks: list[dict[str, str]] = []
    chapter = "未分章"
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name
        if style.startswith("Heading") or re.match(r"^第?[0-9一二三四五六七八九十]+[\.、 ]", text):
            chapter = text[:80]
        blocks.append({"style": style, "chapter": chapter, "text": text})
    for idx, table in enumerate(doc.tables, start=1):
        rows = []
        for row in table.rows[:8]:
            rows.append([cell.text.strip() for cell in row.cells])
        blocks.append({"style": "Table", "chapter": chapter, "text": f"表格{idx}: {rows}"})
    return blocks


def classify_rule(text: str) -> tuple[str, str, str]:
    category = "通用"
    if any(k in text for k in ["战斗", "破绽", "反击", "伤害", "护甲"]):
        category = "战斗"
    elif any(k in text for k in ["情景", "危机", "解密", "线索", "交涉", "潜入", "追逐"]):
        category = "情景"
    elif any(k in text for k in ["休整", "岁月", "成长", "修炼", "名望"]):
        category = "成长/休整/岁月"
    elif any(k in text for k in ["数据库", "生成", "NPC", "剧本"]):
        category = "数据库/生成"
    scene = category
    resources = "、".join([k for k in ["得利", "破绽", "势态", "Pos", "危机", "解密", "名望", "人情", "情报", "把柄", "感悟"] if k in text])
    return category, scene, resources


def make_rule_index(doc_paths: list[Path]) -> list[RuleRecord]:
    records: list[RuleRecord] = []
    counter = 1
    for path in doc_paths:
        source = "玩家规则书" if "玩家" in path.name else "DM规则书"
        blocks = docx_blocks(path)
        for block in blocks:
            text = block["text"]
            if len(text) < 12:
                continue
            if not any(k in text for k in RULE_KEYWORDS) and not block["style"].startswith("Heading"):
                continue
            category, scene, resources = classify_rule(text)
            needs_dm = "是" if any(k in text for k in ["DM", "主持", "裁定", "可选择", "建议"]) else "否"
            ambiguity = "需复核" if any(k in text for k in ["建议", "可选", "通常", "默认", "由DM"]) else "未标记"
            name = text[:36].replace("\n", " ")
            records.append(
                RuleRecord(
                    rule_id=f"R{counter:04d}",
                    rule_name=name,
                    source=source,
                    chapter=block["chapter"],
                    category=category,
                    applicable_scene=scene,
                    input_conditions="待人工细化：由规则文本上下文决定",
                    resolution="待人工细化：已保留原文摘要",
                    affected_resources=resources,
                    linked_database="待匹配",
                    needs_dm_judgment=needs_dm,
                    ambiguity=ambiguity,
                    excerpt=text[:300],
                )
            )
            counter += 1
    return records


def worksheet_rows(ws: Any) -> tuple[list[str], int]:
    rows = list(ws.iter_rows(values_only=True))
    non_empty = [r for r in rows if any(c is not None and str(c).strip() for c in r)]
    header: list[str] = []
    if non_empty:
        best = max(non_empty[:10], key=lambda r: sum(1 for c in r if c is not None and str(c).strip()))
        header = [str(c).strip() if c is not None else "" for c in best]
    return header, max(0, len(non_empty) - 1)


def make_database_index(xlsx_paths: list[Path], rules: list[RuleRecord]) -> list[DatabaseRecord]:
    rule_text = "\n".join(r.excerpt + " " + r.rule_name for r in rules)
    records: list[DatabaseRecord] = []
    for path in xlsx_paths:
        wb = load_workbook(str(path), read_only=True, data_only=True)
        for ws in wb.worksheets:
            header, row_count = worksheet_rows(ws)
            fields = [h for h in header if h]
            pk = fields[0] if fields else "未识别"
            called = "是" if ws.title in rule_text or any(f and f in rule_text for f in fields[:5]) else "未匹配"
            unexplained = [f for f in fields if f and f not in rule_text][:8]
            records.append(
                DatabaseRecord(
                    workbook=path.name,
                    sheet_name=ws.title,
                    primary_key_guess=pk,
                    row_count=row_count,
                    column_count=len(fields),
                    core_fields="、".join(fields[:12]),
                    called_by_rules=called,
                    fields_without_explanation="、".join(unexplained),
                    possibly_unused_fields="待使用率统计复核" if called == "未匹配" else "",
                )
            )
        wb.close()
    return records


def make_term_index(doc_paths: list[Path]) -> list[TermRecord]:
    records: list[TermRecord] = []
    for term in TERMS:
        for path in doc_paths:
            source = "玩家规则书" if "玩家" in path.name else "DM规则书"
            blocks = docx_blocks(path)
            occurrences = 0
            first_chapter = ""
            first_context = ""
            for block in blocks:
                count = block["text"].count(term)
                if count and not first_context:
                    first_chapter = block["chapter"]
                    first_context = block["text"][:220]
                occurrences += count
            records.append(
                TermRecord(
                    term=term,
                    source=source,
                    occurrences=occurrences,
                    first_chapter=first_chapter or "未出现",
                    first_context=first_context,
                    needs_definition_check="是" if occurrences == 0 or "：" not in first_context[:80] else "待复核",
                )
            )
    return records


def main() -> None:
    REPORTS.mkdir(exist_ok=True)
    RAW_LOGS.mkdir(exist_ok=True)

    sources = find_inputs()
    doc_paths = []
    xlsx_paths = []
    seen: set[str] = set()
    for s in sources:
        if not s.found_path or s.status not in {"found", "near_match"}:
            continue
        p = ROOT / s.found_path
        key = str(p.resolve())
        if key in seen:
            continue
        seen.add(key)
        if p.suffix == ".docx":
            doc_paths.append(p)
        elif p.suffix == ".xlsx":
            xlsx_paths.append(p)

    rules = make_rule_index(doc_paths)
    databases = make_database_index(xlsx_paths, rules)
    terms = make_term_index(doc_paths)

    payload = {
        "sources": [asdict(s) for s in sources],
        "rules": [asdict(r) for r in rules],
        "databases": [asdict(d) for d in databases],
        "terms": [asdict(t) for t in terms],
        "summary": {
            "source_count": len(sources),
            "found_docx": len(doc_paths),
            "found_xlsx": len(xlsx_paths),
            "rule_records": len(rules),
            "database_sheets": len(databases),
            "term_records": len(terms),
        },
    }
    (RAW_LOGS / "00_extraction_index.json").write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")

    missing = [s for s in sources if s.status == "missing"]
    summary_lines = [
        "# 输入解析摘要",
        "",
        f"- 已解析规则书：{len(doc_paths)} 个",
        f"- 已解析数据库：{len(xlsx_paths)} 个",
        f"- 规则索引初稿：{len(rules)} 条",
        f"- 数据库工作表索引：{len(databases)} 张表",
        f"- 术语索引：{len(terms)} 条",
        "",
        "## 缺失输入",
    ]
    if missing:
        for item in missing:
            summary_lines.append(f"- {item.expected_name}：{item.note}")
    else:
        summary_lines.append("- 无")
    summary_lines += [
        "",
        "## 注意",
        "",
        "本阶段为自动抽取初稿。规则编号、输入条件、结算方式和数据库调用关系仍需在模拟器迭代中继续细化。",
    ]
    (REPORTS / "00_输入解析摘要.md").write_text("\n".join(summary_lines), encoding="utf-8")


if __name__ == "__main__":
    main()
